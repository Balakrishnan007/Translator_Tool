# -*- coding: utf-8 -*-
"""Parses a .docx file into an ordered list of translatable text segments."""

from docx import Document


def parse_docx(file_path: str) -> list[dict]:
    doc = Document(file_path)
    segments = []
    order = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        segments.append({
            "id": f"p{i}",
            "type": "paragraph",
            "order": order,
            "text": text,
            "style": para.style.name if para.style else None,
        })
        order += 1

    for t_idx, table in enumerate(doc.tables):
        # Grab the header row's text per column, if there is one, so each cell
        # can carry "this is the X column" context. A bare cell like "900"
        # is meaningless in isolation, confirmed by testing that translating
        # it with no context produced a fabricated, wrong result.
        header_texts = []
        if len(table.rows) > 0:
            header_texts = [c.text.strip() for c in table.rows[0].cells]

        for r_idx, row in enumerate(table.rows):
            row_texts = [c.text.strip() for c in row.cells]
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                column_header = header_texts[c_idx] if c_idx < len(header_texts) and r_idx != 0 else None
                row_context = [t for i, t in enumerate(row_texts) if i != c_idx and t]
                segments.append({
                    "id": f"t{t_idx}-r{r_idx}-c{c_idx}",
                    "type": "table_cell",
                    "order": order,
                    "text": text,
                    "table_index": t_idx,
                    "row_index": r_idx,
                    "col_index": c_idx,
                    "column_header": column_header,
                    "row_context": row_context,
                })
                order += 1

    return segments


if __name__ == "__main__":
    import sys
    import json

    path = sys.argv[1] if len(sys.argv) > 1 else r"D:\Rotpunküchen\dataset\01_DE_word_technical_spec.docx"
    result = parse_docx(path)
    print(f"Parsed {len(result)} segments from {path}\n")
    for seg in result[:8]:
        print(json.dumps(seg, ensure_ascii=False))
    print("...")
    for seg in result[-3:]:
        print(json.dumps(seg, ensure_ascii=False))

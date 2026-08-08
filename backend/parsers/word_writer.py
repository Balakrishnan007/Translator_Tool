# -*- coding: utf-8 -*-
"""Writes translated segments back out as a real .docx file (spec section 10).

Takes the same segment list translate_document() already produces: no new
data shape, just a new destination for data that already exists. Handles
segments from any source format: paragraph/table_cell (native to Word),
plus row/table_row (from an Excel or PDF source being exported as Word),
each reconstructed as a real table, not flattened text.
"""

from docx import Document


def write_docx(segments: list[dict], output_path: str, language: str = None) -> str:
    doc = Document()
    if language:
        doc.add_heading(language, level=1)

    segments = sorted(segments, key=lambda s: s["order"])

    table_buffer = {}  # table_index -> {row_index: {col_index: text}}
    current_key = None

    def flush_table():
        nonlocal current_key
        if current_key is None:
            return
        rows = table_buffer.pop(current_key)
        max_col = max(c for row in rows.values() for c in row) if rows else 0
        table = doc.add_table(rows=0, cols=max_col + 1)
        table.style = "Table Grid"
        for row_idx in sorted(rows):
            cells = rows[row_idx]
            row = table.add_row()
            for col_idx in range(max_col + 1):
                row.cells[col_idx].text = cells.get(col_idx, "")
        current_key = None

    flat_row_buffer = []  # for row/table_row segments: (grouping_key, [cell values])
    current_flat_key = None

    def flush_flat_rows():
        nonlocal current_flat_key
        if not flat_row_buffer:
            return
        max_col = max(len(r) for r in flat_row_buffer)
        table = doc.add_table(rows=0, cols=max_col)
        table.style = "Table Grid"
        for cells in flat_row_buffer:
            row = table.add_row()
            for i, val in enumerate(cells):
                row.cells[i].text = val
        flat_row_buffer.clear()
        current_flat_key = None

    for seg in segments:
        text = seg.get("translation") or f"[TRANSLATION FAILED: {seg.get('error', 'unknown error')}]"
        seg_type = seg.get("type")

        if seg_type == "table_cell":
            flush_flat_rows()
            key = seg["table_index"]
            if current_key is not None and current_key != key:
                flush_table()
            current_key = key
            table_buffer.setdefault(key, {}).setdefault(seg["row_index"], {})[seg["col_index"]] = text

        elif seg_type in ("row", "table_row"):
            flush_table()
            key = seg.get("sheet") or seg.get("page")
            if current_flat_key is not None and current_flat_key != key:
                flush_flat_rows()
            current_flat_key = key
            cell_count = len(seg.get("cells") or [])
            parts = [p.strip() for p in text.split(" | ")]
            if cell_count and len(parts) != cell_count:
                # Translation didn't preserve the "|" boundaries. Don't
                # silently misalign columns, keep it as one cell instead.
                parts = [text]
            flat_row_buffer.append(parts)

        else:  # paragraph, text_block, or anything else: plain text
            flush_table()
            flush_flat_rows()
            doc.add_paragraph(text)

    flush_table()
    flush_flat_rows()

    doc.save(output_path)
    return output_path
